import os
import json
import shutil
import threading
import time
from datetime import datetime
from flask import Flask, render_template, jsonify, request, redirect, Response
from lib import tmx_processing

# Version 1.12 - Load concordances on-demand, not in initial payload
# Auto-backup TMX files every 30 minutes

app = Flask(__name__)
PROJECTS_DIR = os.path.join(os.getcwd(), "Projects")


# -------------------------------------------------
# TMX Backup System
# -------------------------------------------------
def backup_tmx(project_path):
    """Create a backup of project.tmx in backups folder, keeping max 3 backups"""
    tmx_file = os.path.join(project_path, "target", "project.tmx")
    
    if not os.path.exists(tmx_file):
        return
    
    # Create backups directory
    backups_dir = os.path.join(project_path, "backups")
    if not os.path.exists(backups_dir):
        os.makedirs(backups_dir)
    
    # Create timestamped backup filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_file = os.path.join(backups_dir, f"project_tmx_backup_{timestamp}.tmx")
    
    # Copy TMX to backup
    shutil.copy2(tmx_file, backup_file)
    
    # Keep only the 3 most recent backups
    backups = sorted(
        [f for f in os.listdir(backups_dir) if f.startswith("project_tmx_backup_") and f.endswith(".tmx")],
        reverse=True
    )
    
    # Delete older backups beyond the 3 most recent
    for old_backup in backups[3:]:
        old_backup_path = os.path.join(backups_dir, old_backup)
        os.remove(old_backup_path)
        print(f"Deleted old backup: {old_backup}")
    
    print(f"Created TMX backup: {backup_file}")


def backup_all_projects():
    """Backup TMX files for all projects"""
    if not os.path.exists(PROJECTS_DIR):
        return
    
    for project_name in os.listdir(PROJECTS_DIR):
        project_path = os.path.join(PROJECTS_DIR, project_name)
        if os.path.isdir(project_path):
            try:
                backup_tmx(project_path)
            except Exception as e:
                print(f"Error backing up {project_name}: {e}")


def backup_timer():
    """Background thread to backup TMX files every 30 minutes"""
    while True:
        time.sleep(1800)  # 30 minutes = 1800 seconds
        backup_all_projects()


# Start backup thread
backup_thread = threading.Thread(target=backup_timer, daemon=True)
backup_thread.start()
print("TMX backup system started (every 30 minutes)")



def list_projects():
    if not os.path.exists(PROJECTS_DIR):
        os.makedirs(PROJECTS_DIR)
    return [
        d for d in os.listdir(PROJECTS_DIR)
        if os.path.isdir(os.path.join(PROJECTS_DIR, d))
    ]


@app.route("/")
def index():
    projects = list_projects()
    if len(projects) == 1:
        project_path = os.path.join(PROJECTS_DIR, projects[0])
        prefs = get_preferences(project_path)
        default_view = prefs.get("default_view", "dashboard")
        
        if default_view == "workbench":
            return redirect(f"/workbench/{projects[0]}")
        else:
            return redirect(f"/dashboard/{projects[0]}")
    return render_template("index.html", projects=projects)


@app.route("/workbench/<project_name>")
def workbench(project_name):
    projects = list_projects()
    if project_name not in projects:
        return jsonify({"error": "Project not found"}), 404
    return render_template("workbench.html", project=project_name, projects=projects)


@app.route("/dashboard/<project_name>")
def dashboard(project_name):
    projects = list_projects()
    if project_name not in projects:
        return jsonify({"error": "Project not found"}), 404
    return render_template("dashboard.html", project=project_name, projects=projects)


# -------------------------------------------------
# User preferences helpers
# -------------------------------------------------
def get_preferences(project_path):
    cache_dir = os.path.join(project_path, "cache")
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)

    prefs_file = os.path.join(cache_dir, "preferences.json")
    if not os.path.exists(prefs_file):
        return {"default_view": "dashboard"}

    try:
        with open(prefs_file, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {"default_view": "dashboard"}


def save_preference(project_path, key, value):
    cache_dir = os.path.join(project_path, "cache")
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)

    prefs_file = os.path.join(cache_dir, "preferences.json")

    prefs = {}
    if os.path.exists(prefs_file):
        try:
            with open(prefs_file, "r", encoding="utf-8") as f:
                prefs = json.load(f)
        except Exception:
            prefs = {}

    prefs[key] = value

    with open(prefs_file, "w", encoding="utf-8") as f:
        json.dump(prefs, f, ensure_ascii=False, indent=2)


# -------------------------------------------------
# Concordance suppressions helpers
# -------------------------------------------------
def get_suppressions(project_path):
    cache_dir = os.path.join(project_path, "cache")
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)

    suppressions_file = os.path.join(cache_dir, "concordance_suppressions.json")
    if not os.path.exists(suppressions_file):
        return {"pos": [], "lemmas": []}

    try:
        with open(suppressions_file, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {"pos": [], "lemmas": []}


def save_suppressions(project_path, suppressions):
    cache_dir = os.path.join(project_path, "cache")
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)

    suppressions_file = os.path.join(cache_dir, "concordance_suppressions.json")

    with open(suppressions_file, "w", encoding="utf-8") as f:
        json.dump(suppressions, f, ensure_ascii=False, indent=2)


def rebuild_working_lemmas(project_path, suppressions):
    """DEPRECATED - kept for backwards compatibility"""
    pass


# -------------------------------------------------
# NAOB overrides helpers
# -------------------------------------------------
def get_naob_overrides(project_path):
    cache_dir = os.path.join(project_path, "cache")
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)

    overrides_file = os.path.join(cache_dir, "naob_overrides.json")
    if not os.path.exists(overrides_file):
        return {}

    try:
        with open(overrides_file, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_naob_override(project_path, lemma_key, url):
    cache_dir = os.path.join(project_path, "cache")
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)

    overrides_file = os.path.join(cache_dir, "naob_overrides.json")

    overrides = {}
    if os.path.exists(overrides_file):
        try:
            with open(overrides_file, "r", encoding="utf-8") as f:
                overrides = json.load(f)
        except Exception:
            overrides = {}

    url = url.strip()
    if not url.startswith(("http://", "https://")):
        url = "https://" + url

    overrides[lemma_key] = url

    with open(overrides_file, "w", encoding="utf-8") as f:
        json.dump(overrides, f, ensure_ascii=False, indent=2)


# -------------------------------------------------
# Meanings helpers
# -------------------------------------------------
def get_meanings(project_path):
    cache_dir = os.path.join(project_path, "cache")
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)

    meanings_file = os.path.join(cache_dir, "meanings.json")
    if not os.path.exists(meanings_file):
        return {}

    try:
        with open(meanings_file, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_meaning(project_path, lemma_key, meaning):
    cache_dir = os.path.join(project_path, "cache")
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)

    meanings_file = os.path.join(cache_dir, "meanings.json")

    meanings = {}
    if os.path.exists(meanings_file):
        try:
            with open(meanings_file, "r", encoding="utf-8") as f:
                meanings = json.load(f)
        except Exception:
            meanings = {}

    meanings[lemma_key] = meaning

    with open(meanings_file, "w", encoding="utf-8") as f:
        json.dump(meanings, f, ensure_ascii=False, indent=2)


@app.route("/project/<project_name>")
def load_project(project_name):
    """V1.12: Load segments with tokens from TMX; lemmas WITHOUT concordances"""
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404

    # V1.11: Load segments WITH tokens from TMX (fast - just XML/JSON parsing)
    segments = tmx_processing.load_segments(project_path, include_tokens=True)
    lemmas = tmx_processing.extract_lemmas(segments, project_path=project_path)
    naob_overrides = get_naob_overrides(project_path)
    meanings = get_meanings(project_path)
    suppressions = get_suppressions(project_path)
    
    source_files = tmx_processing.get_source_files_list(project_path)
    segment_files_map = tmx_processing.get_segment_files_map(project_path)
    
    source_file_filter = request.args.get("source_file")

    seg_list = []
    for s in segments:
        if source_file_filter:
            if segment_files_map.get(s["id"]) != source_file_filter:
                continue
        
        # V1.11: Use tokens from TMX if available
        if s.get("tokens"):
            tokens = s["tokens"]
        else:
            # Fallback: if tokens not in TMX, return empty
            # (This shouldn't happen if concordances were built)
            tokens = []
        
        seg_list.append({
            "id": s["id"],
            "source": s["source"],
            "target": s["target"],
            "tokens": tokens,
            "is_paragraph_end": s["is_paragraph_end"]
        })

    # V1.12: Return lemmas WITHOUT occurrences (loaded on-demand)
    lemma_dict = {}
    for key, data in lemmas.items():
        lemma_dict[key] = {
            "lemma": key.split("|")[0],
            "pos": key.split("|")[1],
            "wordforms": list(data["wordforms"])
            # Occurrences excluded - fetched on-demand via /lemma/<key>/concordances
        }

    return jsonify({
        "segments": seg_list,
        "lemmas": lemma_dict,
        "naob_overrides": naob_overrides,
        "meanings": meanings,
        "suppressions": suppressions,
        "source_files": source_files,
        "segment_files_map": segment_files_map
    })


@app.route("/project/<project_name>/lemma/<lemma_key>/concordances")
def get_lemma_concordances(project_name, lemma_key):
    """V1.12: Fetch concordances for a specific lemma on-demand"""
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    cache_file = tmx_processing.get_cache_file(project_path)
    if not os.path.exists(cache_file):
        return jsonify({"error": "Concordances not built yet"}), 404
    
    try:
        with open(cache_file, "r", encoding="utf-8") as f:
            lemmas = json.load(f)
        
        lemma_data = lemmas.get(lemma_key)
        if not lemma_data:
            return jsonify({"error": "Lemma not found"}), 404
        
        return jsonify({
            "occurrences": lemma_data.get("occurrences", [])
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/project/<project_name>/set_naob_override", methods=["POST"])
def set_naob_override(project_name):
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    lemma_key = data.get("lemma_key")
    url = data.get("url")

    if not lemma_key or not url:
        return jsonify({"error": "Invalid payload"}), 400

    save_naob_override(project_path, lemma_key, url)
    return jsonify({"status": "ok"})


@app.route("/project/<project_name>/set_meaning", methods=["POST"])
def set_meaning(project_name):
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404

    data = request.get_json() or {}
    lemma_key = data.get("lemma_key")
    meaning = data.get("meaning")

    if not lemma_key or not meaning:
        return jsonify({"error": "Invalid payload"}), 400

    save_meaning(project_path, lemma_key, meaning)
    return jsonify({"status": "ok"})


@app.route("/project/<project_name>/update_segment", methods=["POST"])
def update_segment(project_name):
    project_path = os.path.join(PROJECTS_DIR, project_name)
    data = request.get_json()
    seg_id = data.get("segment_id")
    target_text = data.get("target_text", "")
    tmx_processing.update_target_segment(project_path, seg_id, target_text)
    return jsonify({"status": "ok"})


@app.route("/project/<project_name>/lemma_frequencies")
def lemma_frequencies(project_name):
    """Get lemma frequencies - ONLY if master cache already exists"""
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404

    master_file = tmx_processing.get_master_cache_file(project_path)
    if not os.path.exists(master_file):
        return jsonify({"error": "Frequencies not extracted yet. Please extract first."}), 404

    try:
        with open(master_file, "r", encoding="utf-8") as f:
            frequencies = json.load(f)
    except Exception as e:
        return jsonify({"error": f"Failed to load frequencies: {str(e)}"}), 500
    
    suppressions = get_suppressions(project_path)

    pos_groups = {}
    for key, data in frequencies.items():
        lemma = key.split("|")[0]
        pos = key.split("|")[1]
        frequency = data.get("count", len(data.get("occurrences", [])))
        
        if pos not in pos_groups:
            pos_groups[pos] = []
        
        pos_groups[pos].append({
            "lemma_key": key,
            "lemma": lemma,
            "frequency": frequency
        })
    
    for pos in pos_groups:
        pos_groups[pos].sort(key=lambda x: x["frequency"], reverse=True)
    
    return jsonify({
        "pos_groups": pos_groups,
        "suppressions": suppressions
    })


@app.route("/project/<project_name>/extract_frequencies_stream")
def extract_frequencies_stream(project_name):
    """Stream progress updates for initial frequency extraction using SSE"""
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    force = request.args.get('force', 'false').lower() == 'true'
    
    def generate():
        try:
            import time
            
            if force:
                master_file = tmx_processing.get_master_cache_file(project_path)
                if os.path.exists(master_file):
                    os.remove(master_file)
            
            # V1.11: Load segments without tokens (just for iteration)
            segments = tmx_processing.load_segments(project_path, include_tokens=False)
            total = len(segments)
            
            yield f"data: {json.dumps({'status': 'started', 'total_segments': total})}\n\n"
            time.sleep(0.05)
            
            def progress_callback(current, total_items, message, file_name):
                percent = int((current / total_items) * 100) if total_items > 0 else 0
                data = json.dumps({
                    'current': current, 
                    'total': total_items, 
                    'percent': percent, 
                    'message': message, 
                    'file': file_name
                })
                yield f"data: {data}\n\n"
            
            # V1.11: This will tokenize and store to TMX
            for progress_update in tmx_processing.extract_lemma_frequencies_generator(segments, project_path, progress_callback):
                yield progress_update
            
            from datetime import datetime
            final_data = json.dumps({'status': 'complete', 'timestamp': datetime.utcnow().isoformat()})
            yield f"data: {final_data}\n\n"
        except Exception as e:
            import traceback
            error_data = json.dumps({'status': 'error', 'error': str(e), 'traceback': traceback.format_exc()})
            yield f"data: {error_data}\n\n"
    
    response = Response(generate(), mimetype='text/event-stream')
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['X-Accel-Buffering'] = 'no'
    response.headers['Connection'] = 'keep-alive'
    return response


@app.route("/project/<project_name>/file_statistics")
def get_file_statistics(project_name):
    """Get statistics for each source file"""
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    try:
        # Load segments with targets
        segments = tmx_processing.load_segments(project_path, include_tokens=True)
        segment_files_map = tmx_processing.get_segment_files_map(project_path)
        source_files = tmx_processing.get_source_files_list(project_path)
        
        file_stats = []
        for source_file in source_files:
            # Get segments for this file
            file_segment_ids = [seg_id for seg_id, file in segment_files_map.items() if file == source_file]
            file_segments = [s for s in segments if s["id"] in file_segment_ids]
            
            # Count words in source text
            word_count = sum(len(s["source"].split()) for s in file_segments)
            
            # Count translated segments (non-empty target)
            translated_count = sum(1 for s in file_segments if s.get("target", "").strip())
            
            total_segments = len(file_segments)
            percent_translated = round((translated_count / total_segments * 100) if total_segments > 0 else 0, 1)
            
            file_stats.append({
                "filename": source_file,
                "word_count": word_count,
                "total_segments": total_segments,
                "translated_segments": translated_count,
                "percent_translated": percent_translated
            })
        
        return jsonify({"files": file_stats})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/project/<project_name>/export_file/<filename>", methods=["POST"])
def export_file_to_docx(project_name, filename):
    """Export a source file's translations to DOCX with specific formatting"""
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Load segments with targets
        segments = tmx_processing.load_segments(project_path, include_tokens=True)
        segment_files_map = tmx_processing.get_segment_files_map(project_path)
        source_files = tmx_processing.get_source_files_list(project_path)
        
        # Get file number (index in source files list)
        try:
            file_number = source_files.index(filename) + 1
        except ValueError:
            file_number = 1
        
        # Get segments for this file
        file_segment_ids = [seg_id for seg_id, file in segment_files_map.items() if file == filename]
        file_segments = [s for s in segments if s["id"] in file_segment_ids]
        file_segments.sort(key=lambda x: int(x["id"]))
        
        # Create Word document
        doc = Document()
        
        # Set A4 page size and margins (2.54cm = 1 inch)
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A4 height
        section.page_width = Cm(21.0)   # A4 width
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        
        # Create header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = project_name.upper()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.runs[0]
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(8)
        
        # Create footer with page number
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add footer text
        footer_run = footer_para.add_run(f'CHAPTER {file_number} - ')
        footer_run.font.name = 'Times New Roman'
        footer_run.font.size = Pt(8)
        
        # Add page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        # Create a new run for the page number field
        page_run = footer_para.add_run()
        page_run.font.name = 'Times New Roman'
        page_run.font.size = Pt(8)
        page_run._r.append(fldChar1)
        page_run._r.append(instrText)
        page_run._r.append(fldChar2)
        
        # Add chapter heading
        heading = doc.add_paragraph(str(file_number))
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_run = heading.runs[0]
        heading_run.font.name = 'Times New Roman'
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading.paragraph_format.space_after = Pt(12)
        
        # Start first paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 2.0
        
        # Add each segment to current paragraph
        for seg in file_segments:
            target_text = seg.get("target", "").strip()
            if target_text:
                # Add text to current paragraph
                run = p.add_run(target_text + " ")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
                # If this segment ends a paragraph, start a new one
                if seg.get("is_paragraph_end", False):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Cm(0.5)
                    p.paragraph_format.space_after = Pt(12)
                    p.paragraph_format.line_spacing = 2.0
        
        # Save to exports folder
        exports_dir = os.path.join(project_path, "exports")
        if not os.path.exists(exports_dir):
            os.makedirs(exports_dir)
        
        # Generate output filename
        base_name = os.path.splitext(filename)[0]
        output_file = os.path.join(exports_dir, f"{base_name}_translation.docx")
        doc.save(output_file)
        
        return jsonify({
            "status": "ok",
            "filename": f"{base_name}_translation.docx",
            "path": output_file
        })
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "traceback": traceback.format_exc()}), 500


@app.route("/project/<project_name>/suppressions", methods=["GET"])
def get_suppressions_route(project_name):
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    return jsonify(get_suppressions(project_path))


@app.route("/project/<project_name>/suppressions", methods=["POST"])
def save_suppressions_route(project_name):
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    data = request.get_json() or {}
    suppressions = data.get("suppressions", {"pos": [], "lemmas": []})
    
    save_suppressions(project_path, suppressions)
    return jsonify({"status": "ok"})


@app.route("/project/<project_name>/preferences", methods=["GET"])
def get_preferences_route(project_name):
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    return jsonify(get_preferences(project_path))


@app.route("/project/<project_name>/preferences", methods=["POST"])
def save_preference_route(project_name):
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    data = request.get_json() or {}
    key = data.get("key")
    value = data.get("value")
    
    if not key:
        return jsonify({"error": "Invalid payload"}), 400
    
    save_preference(project_path, key, value)
    return jsonify({"status": "ok"})


@app.route("/project/<project_name>/concordance_status")
def concordance_status(project_name):
    """Check if concordances have been built"""
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    cache_file = tmx_processing.get_cache_file(project_path)
    master_file = tmx_processing.get_master_cache_file(project_path)
    
    if not os.path.exists(cache_file):
        return jsonify({
            "built": False,
            "has_master": os.path.exists(master_file)
        })
    
    try:
        with open(cache_file, "r", encoding="utf-8") as f:
            lemmas = json.load(f)
        
        total_lemmas = len(lemmas)
        with_concordances = sum(1 for v in lemmas.values() if v.get("occurrences"))
        
        timestamp = os.path.getmtime(cache_file)
        from datetime import datetime
        dt = datetime.fromtimestamp(timestamp)
        
        return jsonify({
            "built": True,
            "total_lemmas": total_lemmas,
            "with_concordances": with_concordances,
            "timestamp": dt.isoformat()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/project/<project_name>/build_concordances", methods=["POST"])
def build_concordances(project_name):
    """Build full concordances for non-suppressed lemmas only"""
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    try:
        suppressions = get_suppressions(project_path)
        # V1.11: Load segments without tokens (not needed here)
        segments = tmx_processing.load_segments(project_path, include_tokens=False)
        
        tmx_processing.build_working_concordances(project_path, segments, suppressions)
        
        cache_file = tmx_processing.get_cache_file(project_path)
        with open(cache_file, "r", encoding="utf-8") as f:
            lemmas = json.load(f)
        
        total_lemmas = len(lemmas)
        with_concordances = sum(1 for v in lemmas.values() if v.get("occurrences"))
        
        from datetime import datetime
        return jsonify({
            "status": "ok",
            "total_lemmas": total_lemmas,
            "with_concordances": with_concordances,
            "timestamp": datetime.utcnow().isoformat()
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/project/<project_name>/build_concordances_stream")
def build_concordances_stream(project_name):
    """Stream progress updates for concordance building using SSE"""
    project_path = os.path.join(PROJECTS_DIR, project_name)
    if not os.path.exists(project_path):
        return jsonify({"error": "Project not found"}), 404
    
    def generate():
        try:
            import time
            
            suppressions = get_suppressions(project_path)
            # V1.11: Load segments without tokens (not needed here)
            segments = tmx_processing.load_segments(project_path, include_tokens=False)
            total = len(segments)
            
            yield f"data: {json.dumps({'status': 'started', 'total_segments': total})}\n\n"
            time.sleep(0.05)
            
            def progress_callback(current, total_segs, message, file_name):
                percent = int((current / total_segs) * 100) if total_segs > 0 else 0
                data = json.dumps({
                    'current': current, 
                    'total': total_segs, 
                    'percent': percent, 
                    'message': message
                })
                yield f"data: {data}\n\n"
            
            for progress_update in tmx_processing.build_working_concordances_generator(project_path, segments, suppressions, progress_callback):
                yield progress_update
            
            cache_file = tmx_processing.get_cache_file(project_path)
            with open(cache_file, "r", encoding="utf-8") as f:
                lemmas = json.load(f)
            
            total_lemmas = len(lemmas)
            with_concordances = sum(1 for v in lemmas.values() if v.get("occurrences"))
            
            from datetime import datetime
            yield f"data: {json.dumps({'status': 'complete', 'total_lemmas': total_lemmas, 'with_concordances': with_concordances, 'timestamp': datetime.utcnow().isoformat()})}\n\n"
        except Exception as e:
            import traceback
            yield f"data: {json.dumps({'status': 'error', 'error': str(e), 'traceback': traceback.format_exc()})}\n\n"
    
    response = Response(generate(), mimetype='text/event-stream')
    response.headers['Cache-Control'] = 'no-cache'
    response.headers['X-Accel-Buffering'] = 'no'
    response.headers['Connection'] = 'keep-alive'
    return response


if __name__ == "__main__":
    app.run(debug=True)
